"""
续表验证器
验证跨页表格的续表标注格式

规范来源：规范4.3
续表要求：当表格跨页时，需在新页表格上方标注"续表 X-X"或"表 X-X（续）"
"""

import sys
import os
import re

# 添加当前目录到路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from base_validator import BaseValidator, run_validator, parse_args

# 常量定义
FONT_SIZES = {
    '小五': 9,
}

TOLERANCE = {
    'font_pt': 0.5,
}


class TableContinuationValidator(BaseValidator):
    """续表验证器"""
    
    name = "续表验证"
    description = "验证跨页表格的续表标注格式（支持文本框检测）"
    standard_ref = "规范4.3"
    
    # 续表标注的正则模式
    CONTINUATION_PATTERNS = [
        r'续表\s*(\d+)[\-\.](\d+)',           # 续表 1-1
        r'表\s*(\d+)[\-\.](\d+)\s*（续）',     # 表 1-1（续）
        r'表\s*(\d+)[\-\.](\d+)\s*\(续\)',     # 表 1-1(续)
        r'表\s*(\d+)[\-\.](\d+)\s*续',         # 表 1-1 续
        r'Table\s*(\d+)[\-\.](\d+)\s*\(continued\)',  # Table 1-1 (continued)
        r'Table\s*(\d+)[\-\.](\d+)\s*continued',      # Table 1-1 continued
    ]
    
    # 可能需要跨页的表格行数阈值
    LARGE_TABLE_THRESHOLD = 15
    
    def validate(self) -> bool:
        """执行续表验证"""
        
        table_count = len(self.doc.tables)
        table_captions = []
        continuation_captions = []
        large_tables = []  # 可能跨页的大表格
        
        # 检查表格大小
        for i, table in enumerate(self.doc.tables):
            row_count = len(table.rows)
            if row_count >= self.LARGE_TABLE_THRESHOLD:
                large_tables.append((i + 1, row_count))
        
        # 从段落中查找续表标注
        for para in self.doc.paragraphs:
            text = para.text.strip()
            
            # 匹配续表标注
            for pattern in self.CONTINUATION_PATTERNS:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    table_id = f"{match.group(1)}-{match.group(2)}"
                    continuation_captions.append((para, text, table_id, '段落'))
                    break
            else:
                # 匹配普通表题（排除续表）
                match = re.match(r'^表\s*(\d+)[\-\.](\d+)', text)
                if match and '续' not in text:
                    table_id = f"{match.group(1)}-{match.group(2)}"
                    table_captions.append(table_id)
        
        # 从文本框中查找续表标注
        textbox_continuations = self._find_textbox_continuations()
        continuation_captions.extend(textbox_continuations)
        
        # 汇总信息
        if table_count > 0:
            self.add_info(f"检测到 {table_count} 个表格")
        else:
            self.add_info("文档中未检测到表格")
        
        # 报告大表格情况
        if large_tables:
            for table_num, row_count in large_tables:
                self.add_info(f"表格{table_num}有{row_count}行，可能需要跨页")
        
        if table_captions:
            self.add_info(f"检测到 {len(table_captions)} 个表题: {', '.join(['表'+t for t in table_captions[:5]])}" + 
                         (f" 等{len(table_captions)}个" if len(table_captions) > 5 else ""))
        
        if continuation_captions:
            self.add_info(f"检测到 {len(continuation_captions)} 个续表标注:")
            for item in continuation_captions:
                para, text, table_id, source = item
                self.add_info(f"  - 续表{table_id} (来源:{source})")
            
            # 检查续表格式
            for item in continuation_captions:
                para, text, table_id, source = item
                if source == '段落':
                    self._check_continuation_format(para, text, table_id)
        else:
            # 如果有大表格但没有续表标注，给出警告
            if large_tables:
                for table_num, row_count in large_tables:
                    self.add_warning(f"表格{table_num}有{row_count}行可能跨页，但未找到续表标注")
                self.add_warning("跨页表格需在新页表格上方添加续表标注（如'续表 X-X'或'表 X-X（续）'）")
                self.add_info("提示：如果续表标注在文本框中，请确保文本框可被检测")
        
        # 检查续表与表题的对应关系
        self._check_continuation_matching(table_captions, continuation_captions)
        
        return len(self.errors) == 0
    
    def _find_textbox_continuations(self):
        """从文本框中查找续表标注"""
        continuations = []
        
        try:
            # 遍历文档XML查找文本框
            for para in self.doc.paragraphs:
                # 检查段落中的drawing元素（可能包含文本框）
                for run in para.runs:
                    run_xml = run._r.xml
                    
                    # 检查是否包含文本框相关的XML元素
                    if 'w:txbxContent' in run_xml or 'wps:txbx' in run_xml:
                        # 提取文本框内容
                        # 使用正则提取文本内容
                        text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', run_xml)
                        textbox_text = ''.join(text_matches)
                        
                        # 检查是否包含续表标注
                        for pattern in self.CONTINUATION_PATTERNS:
                            match = re.search(pattern, textbox_text, re.IGNORECASE)
                            if match:
                                table_id = f"{match.group(1)}-{match.group(2)}"
                                continuations.append((para, textbox_text, table_id, '文本框'))
                                break
        except Exception as e:
            # 文本框检测失败，静默处理
            pass
        
        return continuations
    
    def _check_continuation_format(self, para, text: str, table_id: str):
        """检查续表格式"""
        
        # 检查居中对齐
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            pass  # 居中正确
        else:
            self.add_error(f"续表{table_id}标注应居中对齐")
        
        # 检查字号（小五号 9pt）
        expected_size = FONT_SIZES['小五']
        
        for run in para.runs:
            if run.font.size:
                actual_size = run.font.size.pt
                if abs(actual_size - expected_size) > TOLERANCE['font_pt']:
                    self.add_warning(f"续表{table_id}标注字号应为小五号(9pt)，当前为{actual_size}pt")
            break
    
    def _check_continuation_matching(self, table_captions: list, continuation_captions: list):
        """检查续表与原表题的对应关系"""
        
        if not continuation_captions:
            return
        
        # 表题编号集合
        table_numbers = set(table_captions)
        
        # 检查续表编号是否有对应的原表题
        for item in continuation_captions:
            para, text, table_id, source = item
            if table_id not in table_numbers:
                self.add_warning(f"续表{table_id}未找到对应的原表题")


if __name__ == '__main__':
    args = parse_args()
    run_validator(TableContinuationValidator, args.doc_path, args.thesis_type)

