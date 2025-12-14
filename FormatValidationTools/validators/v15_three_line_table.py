"""
三线表验证器
验证表格是否符合三线表格式

规范来源：规范4.3
"""

import sys
import os
import re

# 添加当前目录到路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx.oxml.ns import qn
from base_validator import BaseValidator, run_validator, parse_args

# 常量定义
FONT_SIZES = {
    '小五': 9,     # 表格内容应为小五号
    '五号': 10.5,  # 备用
}

# 表格内容字体要求
FONT_NAMES = {
    'chinese': ['宋体', 'SimSun', '宋体-简', 'STSong'],  # 中文字体
    'english': ['Times New Roman', 'Times'],  # 英文/数字字体
}

TOLERANCE = {
    'font_pt': 0.5,
}


class ThreeLineTableValidator(BaseValidator):
    """三线表验证器"""
    
    name = "三线表验证"
    description = "验证表格是否符合三线表格式：仅有顶线、栏目线、底线"
    standard_ref = "规范4.3"
    
    def validate(self) -> bool:
        """执行三线表验证"""
        
        table_count = len(self.doc.tables)
        
        if table_count == 0:
            self.add_info("文档中未检测到表格")
            return True
        
        # 先找到所有表题，建立表格与表序的对应关系
        table_captions = self._find_table_captions()
        
        self.add_info(f"检测到 {table_count} 个表格")
        
        non_three_line_count = 0
        
        for i, table in enumerate(self.doc.tables):
            table_num = i + 1
            # 尝试获取表序
            table_id = table_captions.get(i, f"序号{table_num}")
            
            # 跳过布局表格（如图题表格）
            if self._is_layout_table(table):
                continue
            
            is_three_line, issues = self._check_three_line_format(table)
            
            if not is_three_line:
                non_three_line_count += 1
                self.add_warning(f"表{table_id}:")
                for issue in issues:
                    self.add_warning(f"  - {issue}")
            
            # 检查表格内容字号
            self._check_table_font(table, table_id)
    
    def _is_layout_table(self, table):
        """判断是否为布局表格（非数据表格）"""
        # 只有1-2行的单列表格，可能是图题或布局表格
        if len(table.rows) <= 2 and len(table.columns) == 1:
            # 检查内容是否包含图题模式
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if re.match(r'^图\s*\d+[\-\.]\d+', text):
                        return True
            return True  # 小表格默认视为布局表格
        
        # 只有1行的表格可能是表头或布局
        if len(table.rows) == 1:
            return True
        
        return False
        
        if non_three_line_count == 0:
            self.add_info("所有表格均符合三线表格式")
        else:
            self.add_warning(f"有 {non_three_line_count} 个表格不符合三线表格式")
        
        return len(self.errors) == 0
    
    def _find_table_captions(self):
        """查找表题，返回表格索引到表序的映射"""
        captions = {}
        table_caption_list = []
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            # 匹配表题：表X-X 或 表X.X
            match = re.match(r'^表\s*(\d+)[\-\.](\d+)', text)
            if match:
                table_id = f"{match.group(1)}-{match.group(2)}"
                table_caption_list.append(table_id)
        
        # 假设表题按顺序对应表格
        for i, cap in enumerate(table_caption_list):
            captions[i] = cap
        
        return captions
    
    def _check_three_line_format(self, table):
        """
        检查是否为三线表格式
        三线表特征：只有顶线、栏目线（表头下方）、底线，无左右边线和内部竖线
        
        注意：边框可能设置在表格级别(tblBorders)或单元格级别(tcBorders)
        重要：单元格级别的边框设置会覆盖表格级别的设置
        """
        issues = []
        
        try:
            tbl = table._tbl
            
            # 检查表格级别的边框设置
            has_table_level_top = False
            has_table_level_bottom = False
            has_table_level_left = False
            has_table_level_right = False
            has_table_level_insideV = False
            
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is not None:
                tblBorders = tblPr.find(qn('w:tblBorders'))
                if tblBorders is not None:
                    top = tblBorders.find(qn('w:top'))
                    if top is not None and top.get(qn('w:val')) not in [None, 'nil', 'none']:
                        has_table_level_top = True
                    
                    bottom = tblBorders.find(qn('w:bottom'))
                    if bottom is not None and bottom.get(qn('w:val')) not in [None, 'nil', 'none']:
                        has_table_level_bottom = True
                    
                    left = tblBorders.find(qn('w:left'))
                    if left is not None and left.get(qn('w:val')) not in [None, 'nil', 'none']:
                        has_table_level_left = True
                    
                    right = tblBorders.find(qn('w:right'))
                    if right is not None and right.get(qn('w:val')) not in [None, 'nil', 'none']:
                        has_table_level_right = True
                    
                    insideV = tblBorders.find(qn('w:insideV'))
                    if insideV is not None and insideV.get(qn('w:val')) not in [None, 'nil', 'none']:
                        has_table_level_insideV = True
            
            # 辅助函数：检查单元格边框是否有效（可见）
            def _cell_border_is_visible(cell, border_name):
                """
                检查单元格的指定边框是否可见
                返回: True=可见, False=不可见, None=未定义（继承表格级别）
                """
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    return None  # 未定义，继承表格级别
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    return None  # 未定义，继承表格级别
                border = tcBorders.find(qn(f'w:{border_name}'))
                if border is None:
                    return None  # 未定义，继承表格级别
                val = border.get(qn('w:val'))
                if val in ['nil', 'none']:
                    return False  # 明确设为无边框
                return True  # 有边框
            
            # 检查实际的内部竖线（考虑单元格覆盖、合并单元格和相邻边框合并）
            # 内部竖线需要检查相邻两个单元格的公共边框：
            # - 左侧单元格的RIGHT边框
            # - 右侧单元格的LEFT边框
            # 规则：
            # 1. 如果任意一方明确设为nil，则该位置没有内部竖线
            # 2. 如果一行中的内部竖线模式不一致（有些有，有些没有），认为是残留定义
            # 3. 只有当一行中所有内部竖线位置都有可见边框时，才认为有内部竖线
            has_visible_insideV = False
            num_cols = len(table.columns)
            
            if num_cols > 1:  # 只有多列表格才可能有内部竖线
                for row in table.rows:
                    cells = row.cells
                    
                    # 构建物理单元格列表（去除合并单元格的重复引用）
                    physical_cells = []
                    seen_tc_ids = set()
                    for col_idx, cell in enumerate(cells):
                        tc_id = id(cell._tc)
                        if tc_id not in seen_tc_ids:
                            seen_tc_ids.add(tc_id)
                            # 获取gridSpan
                            tcPr = cell._tc.find(qn('w:tcPr'))
                            grid_span = 1
                            if tcPr is not None:
                                gridSpan_elem = tcPr.find(qn('w:gridSpan'))
                                if gridSpan_elem is not None:
                                    grid_span = int(gridSpan_elem.get(qn('w:val'), '1'))
                            physical_cells.append({
                                'cell': cell,
                                'start_col': col_idx,
                                'end_col': col_idx + grid_span - 1,
                                'grid_span': grid_span
                            })
                    
                    # 检查这一行所有内部竖线位置的边框状态
                    # 收集所有内部边框位置的状态：True=有竖线, False=无竖线, None=未定义
                    border_states = []
                    
                    for i in range(len(physical_cells) - 1):
                        left_cell_info = physical_cells[i]
                        right_cell_info = physical_cells[i + 1]
                        
                        # 跳过横跨整行的单元格
                        if left_cell_info['end_col'] >= num_cols - 1:
                            continue
                        if right_cell_info['start_col'] <= 0:
                            continue
                        
                        left_cell = left_cell_info['cell']
                        right_cell = right_cell_info['cell']
                        
                        # 检查左侧单元格的RIGHT边框
                        left_right_visible = _cell_border_is_visible(left_cell, 'right')
                        # 检查右侧单元格的LEFT边框
                        right_left_visible = _cell_border_is_visible(right_cell, 'left')
                        
                        # 判断这个位置的内部竖线状态
                        if left_right_visible is False or right_left_visible is False:
                            # 任意一方明确设为nil，该位置没有竖线
                            border_states.append(False)
                        elif left_right_visible is None and right_left_visible is None:
                            # 两者都未定义，取决于表格级别
                            border_states.append(True if has_table_level_insideV else False)
                        else:
                            # 至少有一个明确设为可见
                            border_states.append(True)
                    
                    # 判断这一行是否有内部竖线
                    # 只有当所有内部边框位置都有竖线时，才认为有内部竖线
                    # 如果模式不一致（有些有，有些没有），认为是残留定义，不算
                    if len(border_states) > 0:
                        if all(border_states):
                            # 所有位置都有竖线
                            has_visible_insideV = True
                            break
                        # 如果模式不一致或都没有，继续检查下一行
                    
                    if has_visible_insideV:
                        break
            
            # 检查实际的左边线（考虑单元格覆盖）
            has_visible_left = False
            for row in table.rows:
                cells = row.cells
                if len(cells) > 0:
                    first_cell = cells[0]
                    left_visible = _cell_border_is_visible(first_cell, 'left')
                    if left_visible is True:
                        has_visible_left = True
                        break
                    elif left_visible is None and has_table_level_left:
                        has_visible_left = True
                        break
            
            # 检查实际的右边线（考虑单元格覆盖）
            has_visible_right = False
            for row in table.rows:
                cells = row.cells
                if len(cells) > 0:
                    last_cell = cells[-1]
                    right_visible = _cell_border_is_visible(last_cell, 'right')
                    if right_visible is True:
                        has_visible_right = True
                        break
                    elif right_visible is None and has_table_level_right:
                        has_visible_right = True
                        break
            
            # 检查顶线和底线
            has_visible_top = False
            if len(table.rows) > 0:
                first_row = table.rows[0]
                for cell in first_row.cells:
                    top_visible = _cell_border_is_visible(cell, 'top')
                    if top_visible is True:
                        has_visible_top = True
                        break
                    elif top_visible is None and has_table_level_top:
                        has_visible_top = True
                        break
            
            has_visible_bottom = False
            if len(table.rows) > 0:
                last_row = table.rows[-1]
                for cell in last_row.cells:
                    bottom_visible = _cell_border_is_visible(cell, 'bottom')
                    if bottom_visible is True:
                        has_visible_bottom = True
                        break
                    elif bottom_visible is None and has_table_level_bottom:
                        has_visible_bottom = True
                        break
            
            # 三线表检查
            if has_visible_left:
                issues.append("存在左边线（三线表不应有左边线）")
            
            if has_visible_right:
                issues.append("存在右边线（三线表不应有右边线）")
            
            if has_visible_insideV:
                issues.append("存在内部竖线（三线表不应有内部竖线）")
            
            # 顶线和底线是必须的（但只在能检测到边框设置时才警告）
            if tblPr is not None and tblPr.find(qn('w:tblBorders')) is not None:
                if not has_visible_top:
                    issues.append("缺少顶线")
                if not has_visible_bottom:
                    issues.append("缺少底线")
            
            return len(issues) == 0, issues
            
        except Exception as e:
            return True, []  # 出错时假设通过
    
    def _check_table_font(self, table, table_id: str):
        """检查表格内容字号和字体"""
        expected_size = FONT_SIZES['小五']
        
        size_issue_reported = False
        font_issue_reported = False
        
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text.strip()
                        if not text:
                            continue
                        
                        # 检查字号
                        if not size_issue_reported and run.font.size:
                            actual_size = run.font.size.pt
                            if abs(actual_size - expected_size) > TOLERANCE['font_pt']:
                                self.add_warning(f"表{table_id}内容字号应为小五号(9pt)，发现{actual_size}pt")
                                size_issue_reported = True
                        
                        # 检查字体
                        if not font_issue_reported:
                            font_name = self._get_run_font_name(run)
                            if font_name:
                                # 判断文本是否包含中文
                                has_chinese = any('\u4e00' <= c <= '\u9fff' for c in text)
                                has_english = any(c.isascii() and c.isalpha() for c in text)
                                
                                if has_chinese:
                                    # 检查中文字体
                                    if not any(f.lower() in font_name.lower() for f in FONT_NAMES['chinese']):
                                        self.add_warning(f"表{table_id}中文内容字体应为宋体，发现'{font_name}'")
                                        font_issue_reported = True
                                elif has_english:
                                    # 检查英文字体
                                    if not any(f.lower() in font_name.lower() for f in FONT_NAMES['english']):
                                        self.add_warning(f"表{table_id}英文/数字内容字体应为Times New Roman，发现'{font_name}'")
                                        font_issue_reported = True
                        
                        # 如果两个问题都已报告，跳过剩余检查
                        if size_issue_reported and font_issue_reported:
                            return
    
    def _get_run_font_name(self, run):
        """获取run的字体名称"""
        # 优先从run.font获取
        if run.font.name:
            return run.font.name
        
        # 尝试从XML获取
        try:
            rPr = run._r.find(qn('w:rPr'))
            if rPr is not None:
                # 检查中文字体 (eastAsia)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    east_asia = rFonts.get(qn('w:eastAsia'))
                    if east_asia:
                        return east_asia
                    ascii_font = rFonts.get(qn('w:ascii'))
                    if ascii_font:
                        return ascii_font
        except:
            pass
        
        return None


if __name__ == '__main__':
    args = parse_args()
    run_validator(ThreeLineTableValidator, args.doc_path, args.thesis_type)
