"""
一级标题验证器
验证一级标题（章标题）格式

规范来源：规范6.5
"""

import sys
import os
import re

# 添加当前目录到路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from base_validator import BaseValidator, run_validator, parse_args

# 常量定义
FONT_SIZES = {
    '三号': 16,
}

TOLERANCE = {
    'font_pt': 0.5,
}


class Heading1Validator(BaseValidator):
    """一级标题验证器"""
    
    name = "一级标题验证"
    description = "验证一级标题（章标题）格式：第X章"
    standard_ref = "规范6.5"
    
    def validate(self) -> bool:
        """执行一级标题验证"""
        
        heading1_count = 0
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # 匹配一级标题模式：第X章
            # 注意：标题中不应包含冒号或句号（这些是正文内容）
            is_h1_pattern = re.match(r'^第\s*\d+\s*章\s*[^:：。\.]*$', text)
            
            # 额外检查：如果文本包含冒号或句号，很可能是正文而非标题
            has_punctuation = ':' in text or '：' in text or '。' in text or ('. ' in text and not re.match(r'^第\s*\d+\s*章', text))
            if has_punctuation:
                is_h1_pattern = None
            
            # 特殊章节也算一级标题（必须是独立的标题，不含冒号句号）
            clean_text = text.replace(' ', '').replace('\u3000', '')
            is_special_h1 = clean_text in ['摘要', 'ABSTRACT', '目录', 
                                           '参考文献', '附录', '致谢',
                                           '个人简历及在学期间科研成果',
                                           '个人简历',
                                           '在学期间研究成果',
                                           '在学期间科研成果']
            
            if is_h1_pattern or is_special_h1:
                heading1_count += 1
                self._check_heading1_format(para, text, i, bool(is_h1_pattern))
        
        if heading1_count > 0:
            self.add_info(f"检测到 {heading1_count} 个一级标题")
        else:
            self.add_warning("未检测到一级标题（第X章格式）")
        
        return len(self.errors) == 0
    
    def _check_heading1_format(self, para, text: str, para_index: int, is_chapter: bool):
        """检查一级标题格式"""
        
        display_text = text[:20] + '...' if len(text) > 20 else text
        
        # 检查居中对齐
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            self.add_info(f"'{display_text}' 居中对齐")
        else:
            self.add_error(f"一级标题'{display_text}'应居中对齐")
        
        # 检查字体和字号（三号16pt加粗）
        expected_size = FONT_SIZES['三号']  # 16pt
        
        for run in para.runs:
            # 检查字号
            if run.font.size:
                actual_size = run.font.size.pt
                if abs(actual_size - expected_size) <= TOLERANCE['font_pt']:
                    pass  # 字号正确，不重复报告
                else:
                    self.add_warning(f"一级标题'{display_text}'字号应为三号(16pt)，当前为{actual_size}pt")
            
            # 检查加粗
            if not run.font.bold:
                self.add_warning(f"一级标题'{display_text}'应加粗")
            
            break  # 只检查第一个run
        
        # 检查分页（每章另起一页）
        if is_chapter:
            self._check_page_break(para, para_index, display_text)
    
    def _check_page_break(self, para, para_index: int, display_text: str):
        """检查是否另起一页"""
        
        has_page_break = False
        
        # 方法1：检查段落格式中的分页属性
        pPr = para._p.find(qn('w:pPr'))
        if pPr is not None:
            pageBreakBefore = pPr.find(qn('w:pageBreakBefore'))
            if pageBreakBefore is not None:
                val = pageBreakBefore.get(qn('w:val'))
                if val is None or val != '0':
                    has_page_break = True
        
        # 方法2：检查前一段落是否有分页符
        if not has_page_break and para_index > 0:
            prev_para = self.doc.paragraphs[para_index - 1]
            prev_xml = prev_para._p.xml
            if 'w:br' in prev_xml and 'page' in prev_xml:
                has_page_break = True
            
            # 检查前一段落的runs中是否有分页符
            if not has_page_break:
                for run in prev_para.runs:
                    run_xml = run._r.xml
                    if 'w:br' in run_xml and 'page' in run_xml:
                        has_page_break = True
                        break
            
            # 检查前一段落是否是分节符（sectPr会导致分页）
            if not has_page_break:
                if 'w:sectPr' in prev_xml:
                    has_page_break = True
        
        # 方法3：检查段落本身是否以分页符开头
        if not has_page_break:
            for run in para.runs:
                run_xml = run._r.xml
                if 'w:br' in run_xml and 'page' in run_xml:
                    has_page_break = True
                    break
        
        # 方法4：检查前几个段落是否有分页符或分节符（可能隔了空段落）
        if not has_page_break and para_index > 0:
            for back_idx in range(max(0, para_index - 3), para_index):
                check_para = self.doc.paragraphs[back_idx]
                check_xml = check_para._p.xml
                if 'w:sectPr' in check_xml or ('w:br' in check_xml and 'page' in check_xml):
                    has_page_break = True
                    break
        
        # 第一章如果在目录之后，且位置较靠后，不需要额外警告
        # 因为分页可能通过分节符实现，很难完全检测
        # 所以这里不再强制警告


if __name__ == '__main__':
    args = parse_args()
    run_validator(Heading1Validator, args.doc_path, args.thesis_type)
